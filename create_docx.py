from docx import Document
from pathlib import Path
source = Path('d:/netflix/README.md')
dest = Path('d:/netflix/Netflix_Clone_Project_Documentation.docx')
md = source.read_text(encoding='utf-8')
doc = Document()
for line in md.splitlines():
    if line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('  - '):
        para = doc.add_paragraph(line[4:], style='List Bullet 2')
    elif line.startswith('- '):
        para = doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('> '):
        para = doc.add_paragraph(line[2:])
        para.style = 'Intense Quote'
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)
doc.save(dest)
print('saved', dest)
